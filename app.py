from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_paragraph("FORM ‘A’\nMEDIATION APPLICATION FORM\n[REFER RULE 3(1)]")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].bold = True

subtitle = doc.add_paragraph(
    "Mumbai District Legal Services Authority\nCity Civil Court, Mumbai"
)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\nDETAILS OF PARTIES:")

# Create table
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"

# Applicant
row = table.add_row().cells
row[0].text = "1. Name of Applicant"
row[1].text = "{{client_name}}"

row = table.add_row().cells
row[0].text = "Registered Address"
row[1].text = "{{branch_address}}"

row = table.add_row().cells
row[0].text = "Correspondence Address"
row[1].text = "{{branch_address}}"

row = table.add_row().cells
row[0].text = "Telephone / Mobile"
row[1].text = "{{mobile}}"

row = table.add_row().cells
row[0].text = "Email ID"
row[1].text = "info@kslegal.co.in"

# Opposite Party
row = table.add_row().cells
row[0].text = "2. Name of Opposite Party"
row[1].text = "{{customer_name}}"

row = table.add_row().cells
row[0].text = "Registered Address"
row[1].text = "________________________"

row = table.add_row().cells
row[0].text = "Correspondence Address"
row[1].text = "________________________"

# Dispute Section
doc.add_paragraph("\nDETAILS OF DISPUTE:")
doc.add_paragraph(
    "THE COMMERCIAL COURTS (PRE-INSTITUTION MEDIATION AND SETTLEMENT) RULES, 2018"
).runs[0].bold = True

doc.add_paragraph(
    "Nature of disputes as per Section 2(1)(c) of the Commercial Courts Act, 2015:"
)

# Save file
doc.save("Mediation_Application_Form.docx")

print("Word document created successfully.")
